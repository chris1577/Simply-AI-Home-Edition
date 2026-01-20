"""
Document text extraction service.
Extracts text content from various document types for RAG processing.
"""
import os
from pathlib import Path
from typing import Optional
import logging

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Service for extracting text content from documents."""

    # Supported file types
    SUPPORTED_TYPES = {'pdf', 'txt', 'md', 'csv', 'json', 'docx', 'xlsx'}

    @staticmethod
    def extract(file_path: str, file_type: str) -> dict:
        """
        Extract text content from a document.

        Args:
            file_path: Absolute path to the document file
            file_type: Type of document (pdf, txt, docx, etc.)

        Returns:
            dict with keys:
                - text: Full extracted text
                - pages: List of per-page text (for PDFs)
                - metadata: Document metadata
                - error: Error message if extraction failed
        """
        file_type = file_type.lower().strip('.')

        if file_type not in DocumentExtractor.SUPPORTED_TYPES:
            return {
                'text': '',
                'pages': [],
                'metadata': {},
                'error': f'Unsupported file type: {file_type}'
            }

        if not os.path.exists(file_path):
            return {
                'text': '',
                'pages': [],
                'metadata': {},
                'error': f'File not found: {file_path}'
            }

        try:
            if file_type == 'pdf':
                return DocumentExtractor._extract_pdf(file_path)
            elif file_type == 'docx':
                return DocumentExtractor._extract_docx(file_path)
            elif file_type == 'xlsx':
                return DocumentExtractor._extract_xlsx(file_path)
            elif file_type in ('txt', 'md', 'csv', 'json'):
                return DocumentExtractor._extract_text(file_path)
            else:
                return {
                    'text': '',
                    'pages': [],
                    'metadata': {},
                    'error': f'Extractor not implemented for: {file_type}'
                }
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return {
                'text': '',
                'pages': [],
                'metadata': {},
                'error': f'Extraction error: {str(e)}'
            }

    @staticmethod
    def _extract_pdf(file_path: str) -> dict:
        """Extract text from PDF using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            return {
                'text': '',
                'pages': [],
                'metadata': {},
                'error': 'PyMuPDF not installed. Run: pip install PyMuPDF'
            }

        try:
            doc = fitz.open(file_path)
            pages = []
            full_text = []

            metadata = {
                'page_count': len(doc),
                'title': doc.metadata.get('title', ''),
                'author': doc.metadata.get('author', ''),
                'subject': doc.metadata.get('subject', ''),
                'creator': doc.metadata.get('creator', ''),
            }

            for page_num, page in enumerate(doc):
                page_text = page.get_text("text")
                pages.append(page_text)
                full_text.append(page_text)

            doc.close()

            return {
                'text': '\n\n'.join(full_text),
                'pages': pages,
                'metadata': metadata,
                'error': None
            }
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            return {
                'text': '',
                'pages': [],
                'metadata': {},
                'error': f'PDF extraction failed: {str(e)}'
            }

    @staticmethod
    def _extract_docx(file_path: str) -> dict:
        """Extract text from Word documents."""
        try:
            from docx import Document
        except ImportError:
            return {
                'text': '',
                'pages': [],
                'metadata': {},
                'error': 'python-docx not installed. Run: pip install python-docx'
            }

        try:
            doc = Document(file_path)
            paragraphs = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(' | '.join(row_text))
                if table_text:
                    paragraphs.append('\n'.join(table_text))

            # Extract metadata from core properties
            metadata = {}
            try:
                core_props = doc.core_properties
                metadata = {
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'created': str(core_props.created) if core_props.created else '',
                    'modified': str(core_props.modified) if core_props.modified else '',
                }
            except Exception:
                pass

            return {
                'text': '\n\n'.join(paragraphs),
                'pages': [],  # DOCX doesn't have clear page boundaries
                'metadata': metadata,
                'error': None
            }
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            return {
                'text': '',
                'pages': [],
                'metadata': {},
                'error': f'DOCX extraction failed: {str(e)}'
            }

    @staticmethod
    def _extract_xlsx(file_path: str) -> dict:
        """Extract text from Excel spreadsheets."""
        try:
            from openpyxl import load_workbook
        except ImportError:
            return {
                'text': '',
                'pages': [],
                'metadata': {},
                'error': 'openpyxl not installed. Run: pip install openpyxl'
            }

        try:
            wb = load_workbook(file_path, read_only=True, data_only=True)
            sheets_text = []
            pages = []  # Each sheet as a "page"

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_content = [f"## Sheet: {sheet_name}\n"]

                rows = []
                for row in sheet.iter_rows(values_only=True):
                    # Filter out completely empty rows
                    row_values = [str(cell) if cell is not None else '' for cell in row]
                    if any(v.strip() for v in row_values):
                        rows.append(' | '.join(row_values))

                if rows:
                    # Add header separator after first row (assumed to be header)
                    if len(rows) > 1:
                        header_sep = ' | '.join(['---'] * len(rows[0].split(' | ')))
                        rows.insert(1, header_sep)
                    sheet_content.extend(rows)

                sheet_text = '\n'.join(sheet_content)
                sheets_text.append(sheet_text)
                pages.append(sheet_text)

            wb.close()

            metadata = {
                'sheet_count': len(wb.sheetnames),
                'sheet_names': wb.sheetnames,
            }

            return {
                'text': '\n\n'.join(sheets_text),
                'pages': pages,
                'metadata': metadata,
                'error': None
            }
        except Exception as e:
            logger.error(f"XLSX extraction error: {str(e)}")
            return {
                'text': '',
                'pages': [],
                'metadata': {},
                'error': f'XLSX extraction failed: {str(e)}'
            }

    @staticmethod
    def _extract_text(file_path: str) -> dict:
        """Extract text from plain text files (txt, md, csv, json)."""
        try:
            # Try UTF-8 first, then fall back to other encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']

            text = None
            used_encoding = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                return {
                    'text': '',
                    'pages': [],
                    'metadata': {},
                    'error': 'Could not decode file with any supported encoding'
                }

            # Get file extension for metadata
            file_ext = Path(file_path).suffix.lower().strip('.')

            metadata = {
                'encoding': used_encoding,
                'file_type': file_ext,
                'char_count': len(text),
                'line_count': text.count('\n') + 1,
            }

            return {
                'text': text,
                'pages': [],
                'metadata': metadata,
                'error': None
            }
        except Exception as e:
            logger.error(f"Text extraction error: {str(e)}")
            return {
                'text': '',
                'pages': [],
                'metadata': {},
                'error': f'Text extraction failed: {str(e)}'
            }

    @staticmethod
    def get_supported_extensions() -> list:
        """Return list of supported file extensions."""
        return list(DocumentExtractor.SUPPORTED_TYPES)

    @staticmethod
    def is_supported(file_type: str) -> bool:
        """Check if file type is supported for extraction."""
        return file_type.lower().strip('.') in DocumentExtractor.SUPPORTED_TYPES
